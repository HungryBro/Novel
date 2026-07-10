import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_font(run, font_name, size_pt):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    # Set Thai font compatibility
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

def compile_docx():
    doc = Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    episodes_dir = "/Users/dolphin/Desktop/Novel/protocol-arceus/03-episodes"
    font_name = "TH Sarabun New"
    
    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("\n\n\n\n\nผู้ถูกเลือกที่ไม่เคยเลือก\n(Protocol Arceus)\n")
    title_run.bold = True
    set_font(title_run, font_name, 28)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("อาร์คที่ 3: เงาใต้แผนที่ (ตอนที่ 51 - 80)\n\n\n\n")
    set_font(sub_run, font_name, 18)
    
    doc.add_page_break()
    
    # Loop over written episodes of Arc 3 (51 to 80)
    for i in range(51, 81):
        filename = f"episode-{i:02d}.md"
        filepath = os.path.join(episodes_dir, filename)
        
        if not os.path.exists(filepath):
            print(f"Warning: {filename} not found!")
            continue
            
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()
            
        # Separate story from notes
        parts = content.split("---")
        story_part = parts[0].strip()
        
        lines = story_part.split("\n")
        header_line = lines[0].strip()
        
        # Parse title from header line: e.g. # ตอนที่ 51: "ป่าสนที่พึมพำ"
        match = re.match(r"^#\s*ตอนที่\s*(\d+)\s*[:\-]\s*\"?([^\"]+)\"?$", header_line)
        if match:
            ep_num = match.group(1)
            ep_title = match.group(2).strip()
            heading_text = f"ตอนที่ {ep_num} - {ep_title}"
        else:
            heading_text = header_line.replace("#", "").strip()
            
        # Add Chapter Heading
        heading_p = doc.add_paragraph()
        heading_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading_p.paragraph_format.space_before = Pt(18)
        heading_p.paragraph_format.space_after = Pt(12)
        heading_p.paragraph_format.keep_with_next = True
        
        heading_run = heading_p.add_run(heading_text)
        heading_run.bold = True
        set_font(heading_run, font_name, 18)
        
        # Add Story paragraphs
        story_lines = lines[1:]
        for line in story_lines:
            line_str = line.strip()
            if not line_str:
                continue
                
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.first_line_indent = Inches(0.5)
            
            # Clean markdown formatting
            clean_line = line_str.replace("**", "").replace("*", "").replace("`", "")
            
            run = p.add_run(clean_line)
            set_font(run, font_name, 14)
            
        # Add page break after each chapter except the last one
        if i < 80:
            doc.add_page_break()
            
    output_path = "/Users/dolphin/Desktop/Novel/Arc3_Protocol_Arceus.docx"
    doc.save(output_path)
    print(f"Successfully compiled Arc 3 to {output_path}")

if __name__ == "__main__":
    compile_docx()
